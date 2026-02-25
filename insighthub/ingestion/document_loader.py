# ============================================================
#  InsightHub — insighthub/ingestion/document_loader.py
#  Handles: PDF, DOCX, TXT file ingestion
#  Parser:  Unstructured.io (primary), PyMuPDF (fallback)
# ============================================================

import os
import sys
from pathlib import Path
from typing import List, Optional

# Add project root to path
sys.path.append(str(Path(__file__).resolve().parent.parent.parent))

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from insighthub.config.settings import (
    CHUNK_SIZE_DOCUMENT,
    CHUNK_OVERLAP,
    DOCUMENT_EXTENSIONS,
)

# ── Logging setup ─────────────────────────────────────────────
import logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
#  MAIN LOADER CLASS
# ─────────────────────────────────────────────────────────────

class DocumentLoader:
    """
    Loads and chunks PDF, DOCX, and TXT files into LangChain Documents.

    Uses Unstructured.io as the primary parser (extracts text, tables,
    and detects headings). Falls back to PyMuPDF for PDFs and
    python-docx for DOCX files if Unstructured fails.

    Usage:
        loader = DocumentLoader()
        docs   = loader.load("path/to/file.pdf")
        chunks = loader.load_and_chunk("path/to/file.pdf")
    """

    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size    = CHUNK_SIZE_DOCUMENT,
            chunk_overlap = CHUNK_OVERLAP,
            separators    = ["\n\n", "\n", ". ", " ", ""],
            length_function = len,
        )
        logger.info("DocumentLoader initialised (chunk_size=%d, overlap=%d)",
                    CHUNK_SIZE_DOCUMENT, CHUNK_OVERLAP)

    # ── Public methods ────────────────────────────────────────

    def load(self, file_path: str) -> List[Document]:
        """
        Load a single file and return a list of LangChain Documents.
        Each document corresponds to one logical section/page.
        Does NOT chunk — use load_and_chunk() for retrieval.
        """
        path = Path(file_path)
        self._validate_file(path)

        ext = path.suffix.lower()
        logger.info("Loading %s file: %s", ext.upper(), path.name)

        if ext == ".pdf":
            docs = self._load_pdf(path)
        elif ext == ".docx":
            docs = self._load_docx(path)
        elif ext == ".txt":
            docs = self._load_txt(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {DOCUMENT_EXTENSIONS}")

        logger.info("Loaded %d document sections from %s", len(docs), path.name)
        return docs

    def load_and_chunk(self, file_path: str) -> List[Document]:
        """
        Load a file and split into chunks ready for embedding.
        This is the main method used by the ingestion pipeline.

        Returns:
            List of chunked LangChain Documents with full metadata.
        """
        docs   = self.load(file_path)
        chunks = self.splitter.split_documents(docs)

        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"]    = i
            chunk.metadata["total_chunks"] = len(chunks)
            chunk.metadata["source_tab"]  = "document"

        logger.info("Split into %d chunks from %s", len(chunks), Path(file_path).name)
        return chunks

    def load_multiple(self, file_paths: List[str]) -> List[Document]:
        """Load and chunk multiple files at once."""
        all_chunks = []
        for path in file_paths:
            try:
                chunks = self.load_and_chunk(path)
                all_chunks.extend(chunks)
                logger.info("✓ Loaded: %s (%d chunks)", Path(path).name, len(chunks))
            except Exception as e:
                logger.error("✗ Failed to load %s: %s", path, str(e))
        logger.info("Total chunks from %d files: %d", len(file_paths), len(all_chunks))
        return all_chunks

    # ── PDF Loader ────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> List[Document]:
        """Load PDF using Unstructured.io with PyMuPDF fallback."""

        # Try Unstructured.io first
        try:
            return self._load_pdf_unstructured(path)
        except Exception as e:
            logger.warning("Unstructured failed for %s: %s — trying PyMuPDF", path.name, e)
            return self._load_pdf_pymupdf(path)

    def _load_pdf_unstructured(self, path: Path) -> List[Document]:
        """Parse PDF with Unstructured.io — extracts text, tables, headings."""
        from unstructured.partition.pdf import partition_pdf
        from unstructured.documents.elements import Table, Title, NarrativeText

        elements = partition_pdf(
            filename             = str(path),
            extract_images_in_pdf= False,   # skip image extraction for speed
            infer_table_structure= True,    # extract tables as text
            strategy             = "auto",  # auto-selects best parsing method
        )

        docs = []
        current_section = "Introduction"

        for element in elements:
            # Track section headings
            if hasattr(element, "category") and element.category == "Title":
                current_section = str(element)
                continue

            text = str(element).strip()
            if not text or len(text) < 20:  # skip very short fragments
                continue

            has_table = hasattr(element, "category") and element.category == "Table"

            docs.append(Document(
                page_content = text,
                metadata     = {
                    "source"       : str(path),
                    "filename"     : path.name,
                    "source_type"  : "document",
                    "source_tab"   : "document",
                    "file_type"    : "pdf",
                    "section_title": current_section,
                    "has_table"    : has_table,
                    "has_image"    : False,
                    "parser"       : "unstructured",
                }
            ))

        if not docs:
            raise ValueError("Unstructured returned no content")

        return docs

    def _load_pdf_pymupdf(self, path: Path) -> List[Document]:
        """Fallback PDF parser using PyMuPDF — page by page extraction."""
        import fitz  # PyMuPDF

        pdf_doc = fitz.open(str(path))
        docs    = []

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            text = page.get_text("text").strip()

            if not text or len(text) < 20:
                continue

            docs.append(Document(
                page_content = text,
                metadata     = {
                    "source"      : str(path),
                    "filename"    : path.name,
                    "source_type" : "document",
                    "source_tab"  : "document",
                    "file_type"   : "pdf",
                    "page_number" : page_num + 1,
                    "total_pages" : len(pdf_doc),
                    "has_table"   : False,
                    "has_image"   : False,
                    "parser"      : "pymupdf",
                }
            ))

        pdf_doc.close()
        return docs

    # ── DOCX Loader ───────────────────────────────────────────

    def _load_docx(self, path: Path) -> List[Document]:
        """Load DOCX using Unstructured.io with python-docx fallback."""
        try:
            return self._load_docx_unstructured(path)
        except Exception as e:
            logger.warning("Unstructured failed for %s: %s — trying python-docx", path.name, e)
            return self._load_docx_python(path)

    def _load_docx_unstructured(self, path: Path) -> List[Document]:
        """Parse DOCX with Unstructured.io."""
        from unstructured.partition.docx import partition_docx

        elements = partition_docx(filename=str(path))
        docs     = []
        current_section = "Introduction"

        for element in elements:
            if hasattr(element, "category") and element.category == "Title":
                current_section = str(element)
                continue

            text = str(element).strip()
            if not text or len(text) < 20:
                continue

            has_table = hasattr(element, "category") and element.category == "Table"

            docs.append(Document(
                page_content = text,
                metadata     = {
                    "source"       : str(path),
                    "filename"     : path.name,
                    "source_type"  : "document",
                    "source_tab"   : "document",
                    "file_type"    : "docx",
                    "section_title": current_section,
                    "has_table"    : has_table,
                    "has_image"    : False,
                    "parser"       : "unstructured",
                }
            ))

        if not docs:
            raise ValueError("Unstructured returned no content")

        return docs

    def _load_docx_python(self, path: Path) -> List[Document]:
        """Fallback DOCX parser using python-docx — paragraph by paragraph."""
        from docx import Document as DocxDocument

        docx_doc = DocxDocument(str(path))
        docs     = []
        current_section = "Introduction"
        buffer   = []

        for para in docx_doc.paragraphs:
            text  = para.text.strip()
            style = para.style.name if para.style else ""

            if not text:
                continue

            # Detect headings as section markers
            if "Heading" in style:
                # Save buffered content before new section
                if buffer:
                    docs.append(self._make_docx_doc(
                        "\n".join(buffer), path, current_section
                    ))
                    buffer = []
                current_section = text
                continue

            buffer.append(text)

            # Flush buffer every 5 paragraphs to avoid huge chunks
            if len(buffer) >= 5:
                docs.append(self._make_docx_doc(
                    "\n".join(buffer), path, current_section
                ))
                buffer = []

        # Flush remaining buffer
        if buffer:
            docs.append(self._make_docx_doc(
                "\n".join(buffer), path, current_section
            ))

        return docs

    def _make_docx_doc(self, text: str, path: Path, section: str) -> Document:
        """Helper to create a Document with standard DOCX metadata."""
        return Document(
            page_content = text,
            metadata     = {
                "source"       : str(path),
                "filename"     : path.name,
                "source_type"  : "document",
                "source_tab"   : "document",
                "file_type"    : "docx",
                "section_title": section,
                "has_table"    : False,
                "has_image"    : False,
                "parser"       : "python-docx",
            }
        )

    # ── TXT Loader ────────────────────────────────────────────

    def _load_txt(self, path: Path) -> List[Document]:
        """Load plain text file — split by double newlines (paragraphs)."""

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Split by paragraph breaks
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

        docs = []
        for i, para in enumerate(paragraphs):
            if len(para) < 20:
                continue
            docs.append(Document(
                page_content = para,
                metadata     = {
                    "source"      : str(path),
                    "filename"    : path.name,
                    "source_type" : "document",
                    "source_tab"  : "document",
                    "file_type"   : "txt",
                    "paragraph"   : i + 1,
                    "has_table"   : False,
                    "has_image"   : False,
                    "parser"      : "plain-text",
                }
            ))

        return docs

    # ── Validation ────────────────────────────────────────────

    def _validate_file(self, path: Path):
        """Check file exists, is readable, and has supported extension."""
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if not path.is_file():
            raise ValueError(f"Path is not a file: {path}")

        if path.suffix.lower() not in DOCUMENT_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{path.suffix}'. "
                f"Supported types: {DOCUMENT_EXTENSIONS}"
            )

        # Check file is not empty
        if path.stat().st_size == 0:
            raise ValueError(f"File is empty: {path.name}")

        # Check file size — warn if over 50MB
        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > 50:
            logger.warning("Large file detected: %.1fMB — processing may be slow", size_mb)


# ─────────────────────────────────────────────────────────────
#  QUICK TEST — run this file directly to test your loader
#  Command: python insighthub/ingestion/document_loader.py
# ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    print("\n" + "="*55)
    print("  InsightHub — Document Loader Test")
    print("="*55)

    loader = DocumentLoader()

    # ── Test with a TXT file (no extra libraries needed) ─────
    # Creates a sample TXT and loads it to verify everything works

    test_path = Path("test_sample.txt")
    test_path.write_text(
        "Introduction to RAG Systems\n\n"
        "Retrieval-Augmented Generation (RAG) is a technique that combines "
        "information retrieval with language model generation. It was introduced "
        "by Lewis et al. in 2020 and has since become a foundational approach "
        "for knowledge-intensive NLP tasks.\n\n"
        "How RAG Works\n\n"
        "RAG systems first retrieve relevant documents from a knowledge base "
        "using dense vector similarity search. The retrieved documents are then "
        "provided as context to a language model, which generates a grounded "
        "response based on both the query and the retrieved evidence.\n\n"
        "Benefits of RAG\n\n"
        "The main advantage of RAG over pure parametric models is that it "
        "reduces hallucinations by grounding responses in retrieved source "
        "documents. It also allows the knowledge base to be updated without "
        "retraining the language model.",
        encoding="utf-8"
    )

    print("\n[Test 1] Loading TXT file...")
    try:
        chunks = loader.load_and_chunk(str(test_path))
        print(f"  ✓ Loaded {len(chunks)} chunks")
        print(f"  ✓ First chunk preview:")
        print(f"    '{chunks[0].page_content[:100]}...'")
        print(f"  ✓ Metadata: {chunks[0].metadata}")
    except Exception as e:
        print(f"  ✗ Error: {e}")
    finally:
        test_path.unlink()  # clean up test file

    # ── Test with a real PDF if available ────────────────────
    if len(sys.argv) > 1:
        pdf_path = sys.argv[1]
        print(f"\n[Test 2] Loading PDF: {pdf_path}")
        try:
            chunks = loader.load_and_chunk(pdf_path)
            print(f"  ✓ Loaded {len(chunks)} chunks")
            print(f"  ✓ File type: {chunks[0].metadata.get('file_type')}")
            print(f"  ✓ Parser used: {chunks[0].metadata.get('parser')}")
            print(f"  ✓ First chunk: '{chunks[0].page_content[:120]}...'")
        except Exception as e:
            print(f"  ✗ Error: {e}")

    print("\n" + "="*55)
    print("  Document Loader Test Complete!")
    print("  To test with a PDF:")
    print("  python insighthub/ingestion/document_loader.py your_file.pdf")
    print("="*55 + "\n")